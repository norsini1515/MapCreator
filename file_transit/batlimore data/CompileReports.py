# -*- coding: utf-8 -*-
"""
Created on Tue Apr 28 08:08:35 2026
@author: orsin005

CompileReports.py
Read in any number of BLS Data Series reports (SeriesReport-*.xlsx) and
compile them into one long-format table.

Output columns:
    date        - first day of the reference month, as datetime64
    year        - int
    month       - int (1-12); NaN for Annual rows
    value       - float (employees in thousands, or whatever the data type is)
    is_annual   - bool; True for the Dec "Annual" average row entry
    series_info - string combining State / Area / Industry / Data Type
"""

from pathlib import Path
import os
import pandas as pd
import matplotlib.pyplot as plt
import warnings
warnings.filterwarnings("ignore", message="Workbook contains no default style")

working_dir = Path("C:/Users/orsin005/Downloads/batlimore data")

# Map month abbreviations to zero-padded month strings; Annual → "00"
MONTH_MAP = {
    "Jan": "01", "Feb": "02", "Mar": "03", "Apr": "04",
    "May": "05", "Jun": "06", "Jul": "07", "Aug": "08",
    "Sep": "09", "Oct": "10", "Nov": "11", "Dec": "12",
    "Annual": "00",
}
CPI_2024 = {
    2000: 52.6988504,  2001: 53.87958421, 2002: 54.36670309,
    2003: 55.74334341, 2004: 57.28412162,  2005: 59.60852585,
    2006: 62.29297447, 2007: 64.75038334,  2008: 68.7579951,
    2009: 69.29684389, 2010: 70.10487881,  2011: 72.36860498,
    2012: 73.8351505,  2013: 74.9089829,   2014: 76.29658339,
    2015: 76.45500292, 2016: 77.52782931,  2017: 78.98886827,
    2018: 80.4991486,  2019: 81.60946197,  2020: 82.43195161,
    2021: 85.84104252, 2022: 93.67073729,  2023: 97.21484484,
    2024: 100.0,       2025: 103.097123
}

def get_report_info(file_path: Path) -> pd.DataFrame:
    """
    Parse one BLS SeriesReport Excel file and return a long-format DataFrame.

    Parameters
    ----------
    file_path : Path
        Full path to a SeriesReport-*.xlsx file.

    Returns
    -------
    pd.DataFrame
        Long-format data with columns:
        date, year, month, value, is_annual, series_info
    """
    # ------------------------------------------------------------------
    # 1. Read metadata block (rows 5-10, columns A:B in the sheet)
    #    header=4  → row index 4 (0-based) becomes the column header row,
    #    nrows=6   → reads the 6 metadata rows beneath it
    # ------------------------------------------------------------------
    df_info = pd.read_excel(file_path, usecols="A:B", header=4, nrows=6)
    df_info.columns = ["Info_Header", "Info"]
    df_info = df_info.set_index("Info_Header")

    print(
        f"Compiling: {df_info.loc['State:', 'Info']} | "
        f"{df_info.loc['Area:', 'Info']} | "
        f"{df_info.loc['Industry:', 'Info']} | "
        f"{df_info.loc['Data Type:', 'Info']}"
    )

    # ------------------------------------------------------------------
    # 2. Read the data table (header on row 13, i.e. header=12 zero-based)
    #    Columns: Year, Jan, Feb, ..., Dec, Annual
    # ------------------------------------------------------------------
    df_data = pd.read_excel(file_path, usecols="A:N", header=12)

    # Drop any trailing empty rows (Year is NaN when Excel has blank rows)
    df_data = df_data.dropna(subset=["Year"])
    df_data["Year"] = df_data["Year"].astype(int)

    # ------------------------------------------------------------------
    # 3. Melt wide → long
    #    value_vars = the 12 month columns + "Annual"
    # ------------------------------------------------------------------
    month_cols  = list(MONTH_MAP.keys())   # ["Jan", ..., "Dec"]
    value_vars  = month_cols + ["Annual"]

    df_long = df_data.melt(
        id_vars="Year",
        value_vars=value_vars,
        var_name="month_label",
        value_name="value",
    )

    # ------------------------------------------------------------------
    # 4. Derive month string and date
    #    Annual rows get month="00" and date=NaT
    # ------------------------------------------------------------------
    df_long["month"] = df_long["month_label"].map(MONTH_MAP)  # "01"-"12" or "00"
    # Build a proper datetime for non-annual rows (first of the month)
    mask_monthly = df_long["month"] != "00"
    df_long.loc[mask_monthly, "date"] = pd.to_datetime(
        df_long.loc[mask_monthly, "Year"].astype(str)
        + "-"
        + df_long.loc[mask_monthly, "month"]
        + "-01"
    )
    df_long["date"] = pd.to_datetime(df_long["date"])   # NaT for Annual rows

    # ------------------------------------------------------------------
    # 5. Attach metadata columns and tidy up
    # ------------------------------------------------------------------
    df_long["year"]        = df_long["Year"]
    df_long["state"]       = df_info.loc["State:",        "Info"]
    df_long["area"]        = df_info.loc["Area:",         "Info"]
    df_long["supersector"] = df_info.loc["Supersector:",  "Info"]
    df_long["industry"]    = df_info.loc["Industry:",     "Info"]
    df_long["data_type"]   = df_info.loc["Data Type:",    "Info"]

    df_final = (
        df_long[["date", "year", "month", "value",
                 "state", "area", "supersector", "industry", "data_type"]]
        .sort_values(["state", "area", "industry", "date"])
        .reset_index(drop=True)
    )

    return df_final

def plot_annual_by_area(
    df: pd.DataFrame,
    industry: str = None,
    supersector: str = None,
    data_type: str = None,
) -> None:
    """
    Plot annual employment values for all areas in df, optionally filtered
    by industry, supersector, and/or data_type.

    Parameters
    ----------
    df          : compiled master DataFrame from get_report_info
    industry    : filter to a specific industry string (optional)
    supersector : filter to a specific supersector string (optional)
    data_type   : filter to a specific data_type string (optional)
    """
    mask = df["month"] == "00"
    if industry:
        mask &= df["industry"] == industry
    if supersector:
        mask &= df["supersector"] == supersector
    if data_type:
        mask &= df["data_type"] == data_type

    plot_df = df[mask]

    if plot_df.empty:
        print("No data matched the given filters.")
        return

    fig, ax = plt.subplots(figsize=(12, 5))

    for area, grp in plot_df.groupby("area"):
        ax.plot(grp["year"], grp["value"], marker="o", markersize=3, label=area)

    # Build a descriptive title from whatever filters were applied
    filter_parts = [
        p for p in [industry, supersector, data_type] if p
    ]
    title = "Annual Employment — " + (" | ".join(filter_parts) if filter_parts else "All Series")
    ax.set_title(title)
    ax.set_xlabel("Year")
    ax.set_ylabel(plot_df["data_type"].iloc[0])
    ax.legend(fontsize=8)
    ax.grid(True, alpha=0.3)
    fig.tight_layout()
    plt.show()

def compare_areas(
    df: pd.DataFrame,
    area_a: str,
    area_b: str,
    industry: str = None,
    supersector: str = None,
    data_type: str = None,
    base_year: int = None,
) -> None:
    """
    Compare annual trends between two areas across three panels:
      1. Absolute values
      2. Year-over-year % change
      3. Indexed to a base year (both start at 100)
 
    Also prints a summary of key stats.
 
    Parameters
    ----------
    df          : compiled master DataFrame
    area_a      : first area string
    area_b      : second area string
    industry    : filter (optional)
    supersector : filter (optional)
    data_type   : filter (optional)
    base_year   : year to index from; defaults to the first common year
    """
    mask = df["month"] == "00"
    if industry:
        mask &= df["industry"] == industry
    if supersector:
        mask &= df["supersector"] == supersector
    if data_type:
        mask &= df["data_type"] == data_type
 
    annual = df[mask]
 
    a = annual[annual["area"] == area_a][["year", "value"]]
    b = annual[annual["area"] == area_b][["year", "value"]]
 
    if a.empty or b.empty:
        print(f"No data for one or both areas with the given filters.")
        return
 
    # If multiple industries/data_types match, aggregate by summing per year
    a = a.groupby("year")["value"].sum()
    b = b.groupby("year")["value"].sum()
 
    # Align to common years
    common = a.index.intersection(b.index)
    a, b = a.loc[common], b.loc[common]
 
    if base_year is None:
        base_year = int(common[0])
 
    # Derived series
    a_yoy   = a.pct_change(fill_method=None) * 100
    b_yoy   = b.pct_change(fill_method=None) * 100
    a_idx   = (a / a.loc[base_year]) * 100
    b_idx   = (b / b.loc[base_year]) * 100
    diff    = a - b
    ratio   = a / b
 
    # ------------------------------------------------------------------
    # Summary stats
    # ------------------------------------------------------------------
    filter_parts = [p for p in [industry, supersector, data_type] if p]
    label = " | ".join(filter_parts) if filter_parts else "All Series"
    print(f"\n{'='*60}")
    print(f"  Comparison: {area_a}  vs  {area_b}")
    print(f"  Series:     {label}")
    print(f"  Years:      {int(common[0])} – {int(common[-1])}")
    print(f"{'='*60}")
    print(f"{'Metric':<35} {area_a:>10} {area_b:>10}")
    print(f"{'-'*60}")
    print(f"{'Start value':<35} {a.iloc[0]:>10.1f} {b.iloc[0]:>10.1f}")
    print(f"{'End value':<35} {a.iloc[-1]:>10.1f} {b.iloc[-1]:>10.1f}")
    print(f"{'Total growth (%)':<35} {((a.iloc[-1]/a.iloc[0])-1)*100:>10.1f} {((b.iloc[-1]/b.iloc[0])-1)*100:>10.1f}")
    print(f"{'Avg YoY growth (%)':<35} {a_yoy.mean():>10.2f} {b_yoy.mean():>10.2f}")
    print(f"{'Peak value':<35} {a.max():>10.1f} {b.max():>10.1f}")
    print(f"{'Trough value':<35} {a.min():>10.1f} {b.min():>10.1f}")
    print(f"{'Avg absolute gap (A - B)':<35} {diff.mean():>10.1f}")
    print(f"{'Avg ratio (A / B)':<35} {ratio.mean():>10.3f}")
 
    # ------------------------------------------------------------------
    # Plot — 3 panels + difference panel
    # ------------------------------------------------------------------
    fig, axes = plt.subplots(4, 1, figsize=(12, 14), sharex=True)
    fig.suptitle(f"{area_a}  vs  {area_b} | {label}", fontsize=11)
 
    # Panel 1: Absolute
    axes[0].plot(common, a.values, marker="o", markersize=3, label=area_a)
    axes[0].plot(common, b.values, marker="s", markersize=3, linestyle="--", label=area_b)
    axes[0].set_ylabel(data_type or "Value")
    axes[0].set_title("Absolute Values")
    axes[0].legend(fontsize=8)
    axes[0].grid(True, alpha=0.3)
 
    # Panel 2: YoY % change
    axes[1].plot(common, a_yoy.values, marker="o", markersize=3, label=area_a)
    axes[1].plot(common, b_yoy.values, marker="s", markersize=3, linestyle="--", label=area_b)
    axes[1].axhline(0, color="black", linewidth=0.8, linestyle=":")
    axes[1].set_ylabel("YoY % Change")
    axes[1].set_title("Year-over-Year % Change")
    axes[1].legend(fontsize=8)
    axes[1].grid(True, alpha=0.3)
 
    # Panel 3: Indexed
    axes[2].plot(common, a_idx.values, marker="o", markersize=3, label=area_a)
    axes[2].plot(common, b_idx.values, marker="s", markersize=3, linestyle="--", label=area_b)
    axes[2].axhline(100, color="black", linewidth=0.8, linestyle=":")
    axes[2].set_ylabel(f"Index ({base_year} = 100)")
    axes[2].set_title(f"Indexed to {base_year}")
    axes[2].legend(fontsize=8)
    axes[2].grid(True, alpha=0.3)
 
    # Panel 4: Absolute difference (A - B)
    axes[3].bar(common, diff.values, color=["steelblue" if v >= 0 else "tomato" for v in diff.values])
    axes[3].axhline(0, color="black", linewidth=0.8, linestyle=":")
    axes[3].set_ylabel("A − B")
    axes[3].set_title(f"Difference ({area_a} − {area_b})")
    axes[3].set_xlabel("Year")
    axes[3].grid(True, alpha=0.3)
 
    fig.tight_layout()
    plt.show()

def compile_fred_data(
    directory: Path,
    series_names: dict = None,
    real_series: set = None,
) -> pd.DataFrame:
    """
    Read all FRED CSV files in a directory and compile them into one long-format
    DataFrame, optionally deflating monetary series to real 2024 dollars.
 
    Each CSV is expected to have columns: observation_date, <series_id>
 
    Parameters
    ----------
    directory    : Path
        Directory containing FRED CSV files.
    series_names : dict, optional
        Mapping of series_id -> human-readable name, e.g.
        {"MDBALT5POP": "Baltimore Population (5+)"}
        Series not in the dict will have series_name=series_id.
    real_series  : set, optional
        Set of series_ids to deflate using CPI_2024, e.g.
        {"MEHOINUSMDA646N", "MDBALT5URN"}
        Deflated values are stored in a separate `real_value` column.
 
    Returns
    -------
    pd.DataFrame
        Long-format DataFrame with columns:
        observation_date, series, series_name, value
        Monetary series appear twice — once nominal, once with _REAL suffix deflated to 2024$.
    """
    if series_names is None:
        series_names = {}
    if real_series is None:
        real_series = set()
 
    cpi = pd.Series(CPI_2024)
 
    csv_files = [f for f in os.listdir(directory) if f.endswith(".csv")]
    print(f"Found {len(csv_files)} FRED CSV file(s): {csv_files}")
 
    frames = []
    for file in csv_files:
        df = pd.read_csv(directory / file, parse_dates=["observation_date"])
        series_id = [c for c in df.columns if c != "observation_date"][0]
        print(f"  Loaded {file}: {series_id}", end="")
 
        df = df.rename(columns={series_id: "value"})
        df["series"]      = series_id
        df["series_name"] = series_names.get(series_id, series_id)
 
        if series_id in real_series:
            df["year"] = df["observation_date"].dt.year
            real_df = df.copy()
            real_df["value"] = real_df.apply(
                lambda r: (r["value"] / cpi.loc[r["year"]] * 100)
                if r["year"] in cpi.index else float("nan"),
                axis=1,
            )
            real_df["series"]      = series_id + "_REAL"
            real_df["series_name"] = series_names.get(series_id, series_id) + " (Real 2024$)"
            df   = df.drop(columns=["year"])
            real_df = real_df.drop(columns=["year"])
            frames.extend([df, real_df])
            print(" [+ real 2024$ rows added]")
        
        elif series_id == 'MDBALT5POP':
            df['value'] *= 1000
            df['series_name'] = series_names.get(series_id, series_id).replace(" (Thousands)", "")
            frames.append(df)
        else:
            frames.append(df)
            print()
 
    master = (
        pd.concat(frames, axis=0, ignore_index=True)
        .sort_values(["series", "observation_date"])
        .reset_index(drop=True)
    )
 
    return master

def get_fred_snapshot(
    fred_df: pd.DataFrame,
    output_path: Path|None=None,
    anchor_years: list = None,
) -> pd.DataFrame:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if anchor_years is None:
        anchor_years = [2000, 2005, 2010, 2012, 2015, 2020, 2022]

    snapshot = (
        fred_df
        .assign(year=pd.to_datetime(fred_df["observation_date"]).dt.year)
        .query("year in @anchor_years")
        .pivot_table(index="year", columns="series_name", values="value")
        .reset_index()
    )
    
    COLUMNS = {
        "year":                                             "Year",
        
        "Resident Population":                              "Population",
        
        "Civilian Labor Force":                             "Labor Force",
        "Employed Persons":                                 "Employed Persons",
        "Unemployment Rate (%)":                            "Unemp. Rate (%)",
        
        "Median Household Income (Nominal $) (Real 2024$)": "Median HH Income\n(Real 2024$)",
        "All-Transactions House Price Index":               "House Price\nIndex",
    }

    available = [c for c in COLUMNS if c in snapshot.columns]
    table_df  = snapshot[available].rename(columns=COLUMNS)

    if output_path is None:
        return table_df

    def fmt(col, val):
        if pd.isna(val):
            return "—"
        if col == "Year":
            return str(int(val))
        if "Income" in col:
            return f"${val:,.0f}"
        if "Labor" in col or "Population" in col or "Employed" in col:
            return f"{val:,.0f}"
        if "Rate" in col:
            return f"{val:.1f}%"
        return f"{val:,.1f}"

    def set_cell_bg(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def set_cell_borders(cell, i, j, n_rows, n_cols):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right"):
            is_outer = (
                (side == "top"    and i == 0) or
                (side == "left"   and j == 0) or
                (side == "bottom" and i == n_rows - 1) or
                (side == "right"  and j == n_cols - 1)
            )
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"),   "single")
            el.set(qn("w:sz"),    "8" if is_outer else "4")   # 8 = 1pt, 4 = 0.5pt
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
            tcBorders.append(el)
        tcPr.append(tcBorders)

    doc = Document()

    # Page size: landscape letter
    section = doc.sections[0]
    section.page_width  = Inches(11)
    section.page_height = Inches(8.5)
    for margin in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, margin, Inches(1))

    # Caption — font size 10, Aptos
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run("Table 1. Baltimore City Socioeconomic Indicators, Selected Years")
    run.italic         = True
    run.font.size      = Pt(10)
    run.font.name      = "Aptos"
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    headers = list(table_df.columns)
    n_cols  = len(headers)
    table   = doc.add_table(rows=1 + len(table_df), cols=n_cols)
    table.style = "Table Grid"

    # Fixed 1" column widths
    col_width = Inches(1)

    # Disable autofit, enforce fixed layout
    tbl    = table._tbl
    tblPr  = tbl.find(qn("w:tblPr"))
    tblW   = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    "0")
    tblW.set(qn("w:type"), "auto")
    tblPr.append(tblW)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)
    
    tblJc = OxmlElement("w:jc")
    tblJc.set(qn("w:val"), "center")
    tblPr.append(tblJc)

    # Header row
    hdr_row = table.rows[0]
    for j, h in enumerate(headers):
        cell = hdr_row.cells[j]
        cell.width = col_width
        set_cell_bg(cell, "1F4E79")
        set_cell_borders(cell, i=0, j=j, n_rows=1 + len(table_df), n_cols=n_cols)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = para.add_run(h)
        run.bold           = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size      = Pt(9)
        run.font.name      = "Aptos"

    # Data rows
    for i, (_, row) in enumerate(table_df.iterrows()):
        tbl_row = table.rows[i + 1]
        fill    = "F2F2F2" if i % 2 == 0 else "FFFFFF"
        for j, col in enumerate(table_df.columns):
            cell = tbl_row.cells[j]
            cell.width = col_width
            set_cell_bg(cell, fill)
            set_cell_borders(cell, i=i + 1, j=j, n_rows=1 + len(table_df), n_cols=n_cols)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run  = para.add_run(fmt(col, row[col]))
            run.font.size = Pt(9)
            run.font.name = "Aptos"

    # Note — Aptos
    table_width   = Inches(n_cols)
    text_width    = Inches(9)
    left_indent   = (text_width - table_width) / 2
    
    note = doc.add_paragraph()
    note.paragraph_format.left_indent  = int(left_indent)
    note.paragraph_format.right_indent = int(left_indent)
    run = note.add_run(
        "Note: House Price Index indexed to 2000=100. Median household income deflated to "
        "2024 dollars using Baltimore-Columbia-Towson, MD CPI. "
        "Data Sources: Federal Reserve Bank of St. Louis (FRED), Baltimore City Series; "
        "Bureau of Labor Statistics (CPI)."
    )
    run.italic         = True
    run.font.size      = Pt(8)
    run.font.name      = "Aptos"
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.save(output_path)
    print(f"Snapshot saved to {output_path}")

    return table_df
    
if __name__ == "__main__":
    # bls_dir = working_dir /'BLS_BaltimoreEmploymentData'
    # bls_to_compile_files = [
    #     f for f in os.listdir(bls_dir)
    #     if f.startswith("SeriesReport") and f.endswith(".xlsx")
    # ]
    # print(f"{bls_to_compile_files=}")

    # frames = []
    # for file in bls_to_compile_files:
    #     file_path = bls_dir / file
    #     print(f"  Reading {file} ...")
    #     df_series = get_report_info(file_path)
    #     frames.append(df_series)

    # master_df = pd.concat(frames, axis=0, ignore_index=True)
    # master_df.to_csv(working_dir/'BLS_MARYLAND_BALTIMORE_MSA_DATA.csv')
    
    # print(f"\nCompiled {len(bls_to_compile_files)} file(s) → {len(master_df):,} rows")
    # print(master_df.head(10).to_string())
    # print("\nSeries found:")
    # print(master_df[["state", "area", "supersector", "industry", "data_type"]].drop_duplicates().to_string(index=False))

    # combinations = (
    #     master_df[["industry", "data_type"]]
    #     .drop_duplicates()
    #     .sort_values(["industry", "data_type"])
    # )

    # for _, row in combinations.iterrows():
    #     plot_annual_by_area(master_df, industry=row["industry"], data_type=row["data_type"])
        
    # compare_areas(
    #     master_df[master_df['year']!=2026],
    #     area_a="Baltimore-Columbia-Towson, MD",
    #     area_b="Statewide",
    #     industry="Total Private",
    #     data_type="All Employees, In Thousands",
    # )
    
    series_names = {
        'ATNHPIUS24510A': 'All-Transactions House Price Index',
        'GDPALL24510': 'Gross Domestic Product: All Industries (Thousands $)',
        'LAUCN245100000000005A': 'Employed Persons',
        'MDBALT5LFN': 'Civilian Labor Force',
        'MDBALT5POP': 'Resident Population (Thousands)',
        'MDBALT5URN': 'Unemployment Rate (%)',
        'MHIMD24510A052NCEN': 'Median Household Income (Nominal $)',
        'SMS24925810000000001': 'Total Nonfarm Employment (Thousands)'
    }
    fred_dir = working_dir / 'FRED_Data'
    # fred_df = compile_fred_data(fred_dir, series_names=series_names,
    #                             real_series={'GDPALL24510', 'MHIMD24510A052NCEN'})

    # fred_df = fred_df.dropna(axis=0, how='all')
    # fred_df.to_csv(working_dir / 'FRED_BALTIMORE_CITY_DATA.csv', index=False)
    
    fred_df = pd.read_csv(working_dir / 'FRED_BALTIMORE_CITY_DATA.csv')
    
    anchor_years = [1990, 1995, 2000, 2002, 2005, 2010, 2012, 2015, 2020, 2024]
    snapshot_df = get_fred_snapshot(
        fred_df,
        output_path=working_dir / "Baltimore_Snapshot_Table4.docx",
        anchor_years=anchor_years
    )
    # Pivot to wide, filter to anchor years
